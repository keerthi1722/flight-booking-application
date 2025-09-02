from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
import os

# Helper to add code blocks

def add_code_block(doc, code, font_size=9):
    p = doc.add_paragraph()
    run = p.add_run(code)
    font = run.font
    font.name = 'Consolas'
    font.size = Pt(font_size)
    font.color.rgb = RGBColor(50, 50, 50)
    p.style = 'No Spacing'
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# Read code files

def read_file(path):
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    return ''

# Main code files
frontend_app = read_file('client/src/App.js')
frontend_index = read_file('client/src/index.js')
frontend_context = read_file('client/src/context/GeneralContext.jsx')
backend_server = read_file('server/index.js')
backend_schemas = read_file('server/schemas.js')

# Pages/components for flows
login_code = read_file('client/src/components/Login.jsx')
register_code = read_file('client/src/components/Register.jsx')
authenticate_code = read_file('client/src/pages/Authenticate.jsx')
book_flight_code = read_file('client/src/pages/BookFlight.jsx')
bookings_code = read_file('client/src/pages/Bookings.jsx')
new_flight_code = read_file('client/src/pages/NewFlight.jsx')
flight_admin_code = read_file('client/src/pages/FlightAdmin.jsx')
admin_code = read_file('client/src/pages/Admin.jsx')
all_users_code = read_file('client/src/pages/AllUsers.jsx')

# Create document

doc = Document()
doc.add_heading('Flight Finder Rev - Project Report', 0)

# Introduction
intro = (
    "In today's fast-paced world, travel has become an essential part of our lives, "
    "whether it's for business, leisure, or personal reasons. With the advent of technology, "
    "booking flights has become more accessible and convenient than ever before, thanks to flight booking apps. "
    "A flight booking app allows users to search, compare, and book flights easily from the comfort of their devices. "
    "The primary purpose of this app is to streamline the process of planning and booking air travel, providing users with features "
    "that simplify the entire journey, from searching for flights to managing bookings and receiving travel updates."
)
doc.add_heading('Introduction', level=1)
doc.add_paragraph(intro)
doc.add_paragraph(
    "This Flight Finder Rev app is a digital platform designed to revolutionize the way you book flight tickets. "
    "With this app, your flight travel experience is elevated to new heights of convenience and efficiency. "
    "Our user-friendly web app empowers travelers to effortlessly discover, explore, and reserve flight tickets based on their unique preferences. "
    "Whether you're a frequent commuter or an occasional traveler, finding the perfect flight journey has never been easier."
)

# Scenario
scenario = (
    "Scenario:\n"
    "John, a frequent traveler and business professional, needs to book a flight for an upcoming conference. "
    "He prefers using a flight booking app for its convenience and features. John opens the app, enters his travel details, "
    "uses filters to find the best flight, selects his seat, enters payment info, and receives a booking confirmation—all within minutes. "
    "This scenario demonstrates how a flight booking app streamlines the entire travel process for users, offering convenience, customization, and real-time assistance."
)
doc.add_heading('Scenario', level=1)
doc.add_paragraph(scenario)

# Technical Architecture
tech_arch = (
    "In this architecture diagram:\n"
    "- The frontend includes user interface components such as User Authentication, Flight Search, and Booking.\n"
    "- The backend consists of API endpoints for Users, Flights, Admin, and Bookings, including authentication and admin dashboard.\n"
    "- The database stores collections for Users, Flights, and Bookings."
)
doc.add_heading('Technical Architecture', level=1)
doc.add_paragraph(tech_arch)
doc.add_paragraph('[Insert the following Mermaid diagram as an image:]', style='Intense Quote')
doc.add_paragraph('''
graph TD;
  A["Frontend (React)"] -->|"HTTP (REST API)"| B["Backend (Node.js/Express)"];
  B -->|"Mongoose ODM"| C["MongoDB Database"];
  B --> D["Authentication & Authorization"];
  B --> E["Business Logic (Flights, Bookings, Users)"];
  E --> C;
  D --> C;
  A --> F["User (Browser)"];
  F --> A;
''')
doc.add_paragraph('To render this diagram, use an online Mermaid live editor or a diagram tool.')

# ER Diagram
er_desc = (
    "The ER diagram represents the entities and relationships involved in the flight booking system. "
    "It illustrates how users, bookings, flights, and admins are interconnected."
)
doc.add_heading('ER Diagram', level=1)
doc.add_paragraph(er_desc)
doc.add_paragraph('[Insert the following Mermaid ER diagram as an image:]', style='Intense Quote')
doc.add_paragraph('''
erDiagram
  USER {
    string _id PK "User ID"
    string username "Username"
    string email "Email"
    string usertype "User Type (admin/customer/operator)"
    string password "Password Hash"
    string approval "Approval Status"
  }
  FLIGHT {
    string _id PK "Flight ID"
    string flightName "Flight Name"
    string flightId "Flight Number"
    string origin "Origin City"
    string destination "Destination City"
    string departureTime "Departure Time"
    string arrivalTime "Arrival Time"
    number basePrice "Base Price"
    number totalSeats "Total Seats"
  }
  BOOKING {
    string _id PK "Booking ID"
    ObjectId user FK "User Reference"
    ObjectId flight FK "Flight Reference"
    string flightName "Flight Name"
    string flightId "Flight Number"
    string departure "Departure City"
    string destination "Destination City"
    string email "User Email"
    string mobile "User Mobile"
    string seats "Seats"
    array passengers "Passenger List"
    number totalPrice "Total Price"
    date bookingDate "Booking Date"
    date journeyDate "Journey Date"
    string journeyTime "Journey Time"
    string seatClass "Seat Class"
    string bookingStatus "Booking Status"
  }
  USER ||--o{ BOOKING : "books >"
  FLIGHT ||--o{ BOOKING : "< is booked in"
  BOOKING }o--|| USER : "belongs to >"
  BOOKING }o--|| FLIGHT : "< for flight"
''')
doc.add_paragraph('To render this diagram, use an online Mermaid live editor or a diagram tool.')

# Prerequisites
prereq = (
    "To develop a full-stack flight booking app using React JS, Node.js, and MongoDB, you need:\n"
    "- Node.js and npm\n- MongoDB\n- Express.js\n- React.js\n- HTML, CSS, JavaScript\n- Database connectivity (Mongoose)\n- Version control (Git)\n- Development environment (VS Code, Sublime, WebStorm, etc.)"
)
doc.add_heading('Prerequisites', level=1)
doc.add_paragraph(prereq)

# Setup & Installation
setup = (
    "1. Clone the repository\n"
    "2. Install dependencies (npm install in both client and server)\n"
    "3. Start the backend (npm start in server)\n"
    "4. Start the frontend (npm start in client)\n"
    "5. Access the app at http://localhost:3000\n"
    "6. Configure environment variables as needed."
)
doc.add_heading('Setup & Installation', level=1)
doc.add_paragraph(setup)

# Project Structure
proj_struct = (
    "Inside the Flight Finder Rev app directory, we have the following folders:\n"
    "Client directory: Contains the React frontend.\nServer directory: Contains the Node.js/Express backend."
)
doc.add_heading('Project Structure', level=1)
doc.add_paragraph(proj_struct)
doc.add_heading('Client Directory Structure', level=2)
doc.add_paragraph('See client/ for all React source files, components, pages, and assets.')
doc.add_heading('Server Directory Structure', level=2)
doc.add_paragraph('See server/ for all backend code, schemas, and API logic.')

# Application Flow
app_flow = (
    "USER:\n- Create account\n- Search for destination\n- Search for flights\n- Book a flight\n- Make payment\n- Cancel bookings\n\nADMIN:\n- Manage bookings\n- Add new flights\n- Monitor user activity"
)
doc.add_heading('Application Flow', level=1)
doc.add_paragraph(app_flow)

# Project Flow (Milestones)
milestones = (
    "Milestone 1: Project setup and configuration\n"
    "Milestone 2: Implement authentication and user flows\n"
    "Milestone 3: Implement flight search and booking\n"
    "Milestone 4: Admin features and management\n"
    "Milestone 5: Testing and deployment"
)
doc.add_heading('Project Flow (Milestones)', level=1)
doc.add_paragraph(milestones)

# Main Code Sections
code_sections = [
    ('Login Component Code', login_code),
    ('Register Component Code', register_code),
    ('Authenticate Page Code', authenticate_code),
    ('Book Flight Page Code', book_flight_code),
    ('Bookings Page Code', bookings_code),
    ('New Flight Page Code', new_flight_code),
    ('Flight Admin Page Code', flight_admin_code),
    ('Admin Page Code', admin_code),
    ('All Users Page Code', all_users_code),
    ('Frontend Entry (App.js)', frontend_app),
    ('Frontend Entry (index.js)', frontend_index),
    ('Context Provider (GeneralContext.jsx)', frontend_context),
    ('Backend Server (index.js)', backend_server[:3000]),
    ('Schemas (schemas.js)', backend_schemas)
]
for title, code in code_sections:
    doc.add_heading(title, level=2)
    add_code_block(doc, code)

# API Documentation
api_doc = (
    "Authentication:\n- Login: POST /login\n- Register: POST /register\n\nFlight Management: [List endpoints for flights]\nBooking Management: [List endpoints for bookings]\nRequest/Response Examples: [Add examples here]"
)
doc.add_heading('API Documentation', level=1)
doc.add_paragraph(api_doc)

# Deployment Guide
deploy = (
    "Build Steps:\n- Build frontend\n- Prepare backend\n\nEnvironment Configuration: [Add environment variable setup]\nHosting: [Add hosting instructions]"
)
doc.add_heading('Deployment Guide', level=1)
doc.add_paragraph(deploy)

# Changelog
changelog = (
    "[Unreleased]\n- Initial skeleton\n[Date]\n- Description of change"
)
doc.add_heading('Changelog', level=1)
doc.add_paragraph(changelog)

# Roadmap
roadmap = (
    "Planned Features:\n- [ ] Feature 1\n- [ ] Feature 2\nImprovements:\n- [ ] Improvement 1\n- [ ] Improvement 2"
)
doc.add_heading('Roadmap', level=1)
doc.add_paragraph(roadmap)

doc.save('Flight_Finder_Project_Report_Updated.docx')
print('Project report generated: Flight_Finder_Project_Report_Updated.docx') 